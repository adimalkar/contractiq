"""Document processor for extracting structured text, headers, and metadata from PDF and DOCX files."""

import hashlib
import re
from pathlib import Path
from typing import Any

import structlog
from pypdf import PdfReader

from contractiq.pipeline import PageContent, ProcessedDocument, Section

logger = structlog.get_logger(__name__)

# Heuristic patterns for contract headers and sections
SECTION_PATTERNS = [
    re.compile(
        r"^(?:ARTICLE|SECTION|CLAUSE|SCHEDULE|EXHIBIT|APPENDIX)\s+[0-9IVXLC]+[.:]?\s*(.*)$",
        re.IGNORECASE,
    ),
    re.compile(r"^[0-9]+(?:\.[0-9]+)*\s+([A-Z][A-Za-z0-9\s,–—-]+)$"),
    re.compile(r"^[A-Z\s]{4,60}$"),  # Short ALL CAPS lines
]

DATE_PATTERNS = [
    re.compile(
        r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b",
        re.IGNORECASE,
    ),
    re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"),
]

MONEY_PATTERN = re.compile(
    r"\$\s?\d{1,3}(?:,\d{3})*(?:\.\d{2})?(?:\s*(?:million|billion|thousand|USD|dollars))?",
    re.IGNORECASE,
)

PARTY_PATTERNS = [
    re.compile(
        r"(?:between|by and between)\s+([A-Z][A-Za-z0-9\s,.-]+?)\s+(?:and|&)\s+([A-Z][A-Za-z0-9\s,.-]+?)(?:,|\s+effective|\s+dated|\.|$)",
        re.IGNORECASE,
    ),
]

CONTRACT_TYPE_KEYWORDS = [
    "Master Services Agreement",
    "Master Service Agreement",
    "Statement of Work",
    "Non-Disclosure Agreement",
    "Confidentiality Agreement",
    "Service Level Agreement",
    "Vendor Services Agreement",
    "Commercial Lease Agreement",
    "Software License Agreement",
    "Data Processing Agreement",
    "Employment Agreement",
]


class DocumentProcessor:
    """Extracts text, structural sections, and legal metadata from contract files."""

    def compute_hash(self, file_path: Path) -> str:
        """Compute SHA-256 hash of file content for idempotency."""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(65536):
                hasher.update(chunk)
        return hasher.hexdigest()

    def process_file(self, file_path: Path) -> ProcessedDocument:
        """Parse file according to its format."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            return self.process_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self.process_docx(path)
        elif ext in [".txt", ".md"]:
            return self.process_text(path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    def process_pdf(self, file_path: Path) -> ProcessedDocument:
        """Extract structured pages and sections from a PDF file."""
        file_hash = self.compute_hash(file_path)
        reader = PdfReader(str(file_path))
        pages_content: list[PageContent] = []
        full_text_list: list[str] = []

        total_pages = len(reader.pages)
        for page_idx, page in enumerate(reader.pages):
            page_num = page_idx + 1
            raw_page_text = page.extract_text() or ""
            cleaned_text = self._clean_text(raw_page_text)
            full_text_list.append(cleaned_text)

            sections = self._extract_sections(cleaned_text)
            pages_content.append(
                PageContent(
                    page_number=page_num,
                    text=cleaned_text,
                    sections=sections,
                )
            )

        combined_text = "\n\n".join(full_text_list)
        metadata = self._extract_metadata(combined_text, file_path.name)

        return ProcessedDocument(
            filename=file_path.name,
            file_type="pdf",
            file_hash=file_hash,
            page_count=total_pages,
            pages=pages_content,
            metadata=metadata,
            raw_text=combined_text,
        )

    def process_docx(self, file_path: Path) -> ProcessedDocument:
        """Extract structured paragraphs and headings from a DOCX file."""
        file_hash = self.compute_hash(file_path)
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        except Exception as e:
            logger.warning("python-docx parsing fallback to binary read", error=str(e))
            paragraphs = []

        combined_text = "\n\n".join(paragraphs)
        metadata = self._extract_metadata(combined_text, file_path.name)

        # Synthesize pages (~2500 characters per page)
        chars_per_page = 2500
        page_texts = [
            combined_text[i : i + chars_per_page]
            for i in range(0, max(len(combined_text), 1), chars_per_page)
        ]
        pages_content = [
            PageContent(
                page_number=idx + 1,
                text=text,
                sections=self._extract_sections(text),
            )
            for idx, text in enumerate(page_texts)
        ]

        return ProcessedDocument(
            filename=file_path.name,
            file_type="docx",
            file_hash=file_hash,
            page_count=len(pages_content),
            pages=pages_content,
            metadata=metadata,
            raw_text=combined_text,
        )

    def process_text(self, file_path: Path) -> ProcessedDocument:
        """Process plain text or markdown file."""
        file_hash = self.compute_hash(file_path)
        with open(file_path, encoding="utf-8", errors="replace") as f:
            text = f.read()

        cleaned_text = self._clean_text(text)
        metadata = self._extract_metadata(cleaned_text, file_path.name)
        sections = self._extract_sections(cleaned_text)

        return ProcessedDocument(
            filename=file_path.name,
            file_type=file_path.suffix.lstrip("."),
            file_hash=file_hash,
            page_count=1,
            pages=[PageContent(page_number=1, text=cleaned_text, sections=sections)],
            metadata=metadata,
            raw_text=cleaned_text,
        )

    def _clean_text(self, text: str) -> str:
        """Normalize line breaks and clean whitespace while preserving structure."""
        # Replace multiple spaces with a single space, but keep paragraph breaks
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        # Remove empty lines excess
        cleaned = "\n".join(lines)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _extract_sections(self, text: str) -> list[Section]:
        """Detect section headers and split content accordingly."""
        lines = text.splitlines()
        sections: list[Section] = []
        current_header: str | None = None
        current_lines: list[str] = []
        current_offset = 0

        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                current_lines.append("")
                continue

            is_header = False
            for pattern in SECTION_PATTERNS:
                if pattern.match(trimmed):
                    is_header = True
                    break

            if is_header and len(trimmed) < 120:
                # Save previous section if it has text
                if current_lines:
                    sec_text = "\n".join(current_lines).strip()
                    if sec_text:
                        sections.append(
                            Section(
                                header=current_header,
                                text=sec_text,
                                start_offset=current_offset,
                            )
                        )
                current_header = trimmed
                current_lines = [trimmed]
            else:
                current_lines.append(line)

        if current_lines:
            sec_text = "\n".join(current_lines).strip()
            if sec_text:
                sections.append(
                    Section(
                        header=current_header,
                        text=sec_text,
                        start_offset=current_offset,
                    )
                )

        if not sections and text.strip():
            sections.append(Section(header=None, text=text.strip(), start_offset=0))

        return sections

    def _extract_metadata(self, text: str, filename: str) -> dict[str, Any]:
        """Extract key legal and contractual metadata via regex and keywords."""
        metadata: dict[str, Any] = {
            "dates_found": [],
            "amounts_found": [],
            "parties": [],
            "contract_type": "Unknown Agreement",
        }

        # Contract type detection
        for kw in CONTRACT_TYPE_KEYWORDS:
            if kw.lower() in text[:2000].lower() or kw.lower() in filename.lower():
                metadata["contract_type"] = kw
                break

        # Dates extraction
        for d_pat in DATE_PATTERNS:
            for match in d_pat.finditer(text):
                d_str = match.group(0).strip()
                if d_str not in metadata["dates_found"]:
                    metadata["dates_found"].append(d_str)
        metadata["dates_found"] = metadata["dates_found"][:5]  # Keep top 5 dates

        # Financial amounts
        amounts = set(MONEY_PATTERN.findall(text))
        metadata["amounts_found"] = sorted(list(amounts))[:8]

        # Parties extraction
        for p_pat in PARTY_PATTERNS:
            match = p_pat.search(text[:2500])
            if match:
                party_a = match.group(1).strip(" ,.\n")
                party_b = match.group(2).strip(" ,.\n")
                if party_a and party_b:
                    metadata["parties"] = [party_a, party_b]
                break

        return metadata
